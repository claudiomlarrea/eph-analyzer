"""Exportación de resultados a Excel y Word."""

from __future__ import annotations

import io
import json
import sys
from pathlib import Path
from typing import Any

import pandas as pd

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.export_excel import preparar_df_para_excel, preparar_metadatos_para_excel

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:  # pragma: no cover
    Document = None


def _banda_0_1(valor: float) -> str:
    if pd.isna(valor):
        return "N/D"
    if valor < 0.33:
        return "Bajo"
    if valor < 0.66:
        return "Medio"
    return "Alto"


def _interpretar_indice(nombre: str, valor: float, *, invertido: bool = False) -> str:
    banda = _banda_0_1(valor)
    if banda == "N/D":
        return f"{nombre}: sin datos para interpretar."
    if invertido:
        lectura = {"Bajo": "desfavorable", "Medio": "intermedio", "Alto": "favorable"}[banda]
        return f"{nombre}: nivel {banda.lower()} ({valor:.3f}), desempeño {lectura}."
    lectura = {"Bajo": "favorable", "Medio": "intermedia", "Alto": "desfavorable"}[banda]
    return f"{nombre}: nivel {banda.lower()} ({valor:.3f}), situación {lectura}."


def resumen_interpretacion_indices(desc: pd.DataFrame) -> list[dict[str, str | float]]:
    """Arma una tabla de interpretación para los índices principales (escala 0-1)."""
    if desc is None or desc.empty or "anio" not in desc.columns:
        return []
    ref = desc.sort_values("anio").iloc[-1]
    anio = int(ref["anio"])
    filas: list[dict[str, str | float]] = []

    def add(indicador: str, col: str, invertido: bool = False) -> None:
        if col not in ref.index:
            return
        valor = pd.to_numeric(ref[col], errors="coerce")
        if pd.isna(valor):
            return
        filas.append(
            {
                "indicador": indicador,
                "anio_referencia": anio,
                "valor": round(float(valor), 4),
                "nivel": _banda_0_1(float(valor)),
                "interpretacion": _interpretar_indice(indicador, float(valor), invertido=invertido),
            }
        )

    add("Exclusión digital", "idx_exclusion_digital", invertido=False)
    add("Movilidad social (proxy)", "score_movilidad_proxy", invertido=True)
    add("Vulnerabilidad social", "vulnerabilidad_social", invertido=False)
    return filas


def _escribir_excel(resultado: dict, destino) -> None:
    tablas: dict[str, pd.DataFrame] = resultado.get("tablas", {})
    with pd.ExcelWriter(destino, engine="openpyxl") as writer:
        meta = resultado.get("meta", {})
        if meta:
            preparar_metadatos_para_excel(meta).to_excel(writer, sheet_name="metadatos", index=False)
        for nombre, hoja in [
            ("descriptivos", "descriptivos_anuales"),
            ("frecuencias", "frecuencias"),
            ("correlaciones", "correlaciones"),
            ("logistica_coef", "logistica_coeficientes"),
            ("logistica_or", "logistica_odds_ratios"),
            ("shap", "shap_importancia"),
            ("cluster_tamanos", "cluster_tamanos"),
            ("cluster_perfiles", "cluster_perfiles"),
        ]:
            df = tablas.get(hoja)
            if df is not None and not df.empty:
                preparar_df_para_excel(df, incluir_codigo=True).to_excel(
                    writer, sheet_name=nombre[:31], index=False
                )
        modelos = resultado.get("modelos", {})
        if modelos:
            pd.DataFrame([modelos]).to_excel(writer, sheet_name="resumen_modelos", index=False)


def exportar_excel_bytes(resultado: dict) -> bytes:
    buf = io.BytesIO()
    _escribir_excel(resultado, buf)
    buf.seek(0)
    return buf.getvalue()


def exportar_excel(resultado: dict, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    _escribir_excel(resultado, path)
    return path


def _agregar_tabla(doc: Any, titulo: str, df: pd.DataFrame, max_filas: int = 30) -> None:
    doc.add_heading(titulo, level=2)
    if df is None or df.empty:
        doc.add_paragraph("Sin datos disponibles.")
        return

    sub = df.head(max_filas)
    table = doc.add_table(rows=1, cols=len(sub.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    from src.etiquetador import nombre_variable

    for i, col in enumerate(sub.columns):
        hdr[i].text = nombre_variable(str(col))

    for _, row in sub.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(round(val, 4) if isinstance(val, float) else val)


def exportar_word(
    resultado: dict,
    path: Path,
    *,
    titulo: str,
    periodo: str,
    ambito: str,
) -> Path:
    """Genera informe Word con resumen ejecutivo y tablas principales."""
    if Document is None:
        raise ImportError("Instale python-docx: pip install python-docx")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading(titulo, level=0)
    doc.add_paragraph(f"Ámbito: {ambito}")
    doc.add_paragraph(f"Período: {periodo}")
    doc.add_paragraph(f"Fuente: INDEC — EPH (módulos hogar, individuo y TIC)")

    meta = resultado.get("meta", {})
    doc.add_heading("Resumen ejecutivo", level=1)
    doc.add_paragraph(
        f"Se analizaron {meta.get('registros', '—'):,} registros individuales (15+ años) "
        f"con variables proxy de exclusión digital y movilidad social."
    )

    corr = resultado.get("correlacion_destacada")
    if corr is not None:
        doc.add_paragraph(
            f"Correlación Pearson entre exclusión digital y score de movilidad proxy: {corr:.3f}"
        )

    tablas = resultado.get("tablas", {})
    guia = resumen_interpretacion_indices(tablas.get("descriptivos_anuales"))
    if guia:
        doc.add_heading("Guía de interpretación de índices (escala 0 a 1)", level=1)
        doc.add_paragraph("Rangos sugeridos: Bajo < 0,33 · Medio 0,33-0,66 · Alto > 0,66.")
        for fila in guia:
            doc.add_paragraph(f"- {fila['interpretacion']}")

    _agregar_tabla(doc, "Indicadores anuales", tablas.get("descriptivos_anuales"))
    _agregar_tabla(doc, "Frecuencias ponderadas", tablas.get("frecuencias"))
    _agregar_tabla(doc, "Correlaciones", tablas.get("correlaciones"))

    modelos = resultado.get("modelos", {})
    if modelos:
        doc.add_heading("Modelos estadísticos", level=1)
        for clave, valor in modelos.items():
            if isinstance(valor, (dict, list)):
                texto = json.dumps(valor, ensure_ascii=False, indent=2)
            else:
                texto = str(valor)
            doc.add_paragraph(f"{clave}: {texto[:2000]}")

    _agregar_tabla(doc, "Regresión logística — coeficientes", tablas.get("logistica_coeficientes"))
    _agregar_tabla(doc, "Regresión logística — odds ratios", tablas.get("logistica_odds_ratios"))
    _agregar_tabla(doc, "Importancia SHAP (peso relativo %)", tablas.get("shap_importancia"))
    _agregar_tabla(doc, "Clústeres — tamaño (%)", tablas.get("cluster_tamanos"))
    _agregar_tabla(doc, "Clústeres — perfiles medios", tablas.get("cluster_perfiles"))

    grafico = resultado.get("grafico_shap")
    if grafico and Path(grafico).exists():
        doc.add_heading("Gráfico SHAP", level=2)
        doc.add_picture(str(grafico), width=Inches(6))

    doc.add_paragraph(
        "Nota metodológica: índice de exclusión digital compuesto (acceso hogar + uso individual TIC); "
        "movilidad social operada como proxies educativos, laborales y de ingreso. "
        "Fuente primaria: INDEC (www.indec.gob.ar)."
    )

    doc.save(path)
    return path


def exportar_word_bytes(
    resultado: dict,
    *,
    titulo: str,
    periodo: str,
    ambito: str,
) -> bytes:
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        p = Path(f.name)
    try:
        exportar_word(resultado, p, titulo=titulo, periodo=periodo, ambito=ambito)
        return p.read_bytes()
    finally:
        p.unlink(missing_ok=True)
