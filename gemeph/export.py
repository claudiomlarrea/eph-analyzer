"""Exportación Excel, Word y JSON para GEMEPH."""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

import pandas as pd

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


def _meta_df(meta: dict[str, Any]) -> pd.DataFrame:
    return pd.DataFrame([{"campo": k, "valor": str(v)} for k, v in meta.items()])


def export_catalog_excel_bytes(cat_df: pd.DataFrame, meta: dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        _meta_df(meta).to_excel(writer, sheet_name="metadatos", index=False)
        cat_df.to_excel(writer, sheet_name="catalogo_territorial", index=False)
        aglo = cat_df.loc[cat_df["tipo"] == "aglomerado"].copy()
        if not aglo.empty and "idx_exclusion_digital" in aglo.columns:
            aglo.sort_values("idx_exclusion_digital", ascending=False).head(10).to_excel(
                writer, sheet_name="top_exclusion_digital", index=False
            )
    buf.seek(0)
    return buf.getvalue()


def export_baseline_excel_bytes(baseline: dict[str, Any], meta: dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    kpis = dict(baseline.get("kpis", {}))
    brechas = kpis.pop("brechas", {}) if isinstance(kpis.get("brechas"), dict) else {}
    kpis_flat = dict(kpis)
    if brechas:
        for k, v in brechas.items():
            kpis_flat[f"brecha_{k}"] = v

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        _meta_df({**meta, "territorio": baseline.get("territorio_nombre")}).to_excel(
            writer, sheet_name="metadatos", index=False
        )
        pd.DataFrame([kpis_flat]).to_excel(writer, sheet_name="kpis", index=False)
        evo = baseline.get("evolucion") or []
        if evo:
            pd.DataFrame(evo).to_excel(writer, sheet_name="evolucion", index=False)
        perfiles = baseline.get("perfiles") or []
        if perfiles:
            pd.DataFrame(perfiles).to_excel(writer, sheet_name="perfiles", index=False)
    buf.seek(0)
    return buf.getvalue()


def export_scenario_excel_bytes(scenario_result: dict[str, Any], meta: dict[str, Any]) -> bytes:
    from .scenario import compare_rows

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        _meta_df(meta).to_excel(writer, sheet_name="metadatos", index=False)
        compare_rows(scenario_result).to_excel(writer, sheet_name="baseline_vs_escenario", index=False)
        pd.DataFrame([scenario_result.get("targets", {})]).to_excel(writer, sheet_name="palancas", index=False)
        pd.DataFrame([scenario_result.get("modelo", {})]).to_excel(writer, sheet_name="modelo_logistico", index=False)
        pd.DataFrame([scenario_result.get("baseline_kpis", {})]).to_excel(
            writer, sheet_name="kpis_baseline", index=False
        )
        pd.DataFrame([scenario_result.get("scenario_kpis", {})]).to_excel(
            writer, sheet_name="kpis_escenario", index=False
        )
    buf.seek(0)
    return buf.getvalue()


def export_catalog_json_bytes(catalog: dict[str, Any]) -> bytes:
    return json.dumps(catalog, ensure_ascii=False, indent=2).encode("utf-8")


def export_scenario_json_bytes(scenario_result: dict[str, Any], meta: dict[str, Any]) -> bytes:
    payload = {"meta": meta, **scenario_result}
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def export_word_bytes(
    *,
    titulo: str,
    periodo: str,
    cat_df: pd.DataFrame,
    baseline: dict[str, Any] | None = None,
    scenario: dict[str, Any] | None = None,
) -> bytes:
    if Document is None:
        raise ImportError("Instale python-docx")

    doc = Document()
    doc.add_heading(titulo, level=0)
    doc.add_paragraph(f"Período: {periodo}")
    doc.add_paragraph("Fuente: INDEC — EPH (hogar + individuo). GEMEPH — gemelo sociodemográfico territorial.")
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Resumen nacional y aglomerados", level=1)
    nacional = cat_df.loc[cat_df["territorio_id"] == "nacional"]
    if not nacional.empty:
        row = nacional.iloc[0]
        doc.add_paragraph(
            f"Argentina (31 aglomerados): ocupación {row.get('pct_ocupado', '—')}%, "
            f"educación superior {row.get('pct_superior', '—')}%, "
            f"exclusión digital {row.get('idx_exclusion_digital', '—')} (índice)."
        )

    aglo = cat_df.loc[cat_df["tipo"] == "aglomerado"].copy()
    if "idx_exclusion_digital" in aglo.columns:
        top = aglo.nlargest(5, "idx_exclusion_digital")[["territorio_nombre", "idx_exclusion_digital"]]
        doc.add_heading("Mayor exclusión digital (top 5)", level=2)
        for _, r in top.iterrows():
            doc.add_paragraph(f"- {r['territorio_nombre']}: {r['idx_exclusion_digital']:.4f}")

    if baseline:
        doc.add_heading(f"Estado — {baseline.get('territorio_nombre', '')}", level=1)
        kpis = baseline.get("kpis", {})
        for k, v in kpis.items():
            if k == "brechas" and isinstance(v, dict):
                for bk, bv in v.items():
                    doc.add_paragraph(f"Brecha {bk}: {bv}")
            elif k != "brechas":
                doc.add_paragraph(f"{k}: {v}")

    if scenario:
        doc.add_heading("Escenario contrafactual", level=1)
        from .scenario import compare_rows

        cmp = compare_rows(scenario)
        for _, r in cmp.iterrows():
            doc.add_paragraph(
                f"{r['Indicador']}: baseline {r['Baseline']} → escenario {r['Escenario']} "
                f"(Δ {r.get('Cambio', '—')})"
            )
        modelo = scenario.get("modelo", {})
        if modelo.get("pct_exclusion_predicho_base") is not None:
            doc.add_paragraph(
                f"Probabilidad predictiva exclusión alta: {modelo['pct_exclusion_predicho_base']}% → "
                f"{modelo.get('pct_exclusion_predicho_escenario', '—')}%"
            )

    doc.add_paragraph(
        "Nota: escenarios y mapas son análisis contrafactuales; no constituyen proyección oficial INDEC."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
